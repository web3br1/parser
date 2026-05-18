from pathlib import Path

from parsers.base import (
    MAX_EXTRACTED_CHARS,
    BaseParser,
    ExtractedPage,
    ExtractionError,
    ExtractionResult,
    sanitize_text,
    truncate_to_limit,
)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DOCXParser(BaseParser):
    def extract(self, path: Path) -> ExtractionResult:
        try:
            from docx import Document

            document = Document(str(path))
            blocks: list[str] = []

            for paragraph in document.paragraphs:
                text = sanitize_text(paragraph.text)
                if not text:
                    continue
                style_name = paragraph.style.name.lower() if paragraph.style else ""
                if style_name.startswith("heading 1"):
                    blocks.append(f"## {text}")
                elif style_name.startswith("heading"):
                    blocks.append(f"### {text}")
                else:
                    blocks.append(text)

            for table in document.tables:
                for row in table.rows:
                    cells = [sanitize_text(cell.text) for cell in row.cells]
                    blocks.append("| " + " | ".join(cells) + " |")

            text = sanitize_text("\n\n".join(blocks))
            warnings: list[str] = []
            text, truncated = truncate_to_limit(text, 0)
            if truncated:
                warnings.append("text_limit_reached")

            char_count = len(text)
            page = ExtractedPage(
                page_number=0,
                text=text,
                char_count=char_count,
                is_empty=char_count == 0,
            )
            return ExtractionResult(
                mime_type=DOCX_MIME,
                pages=[page],
                total_chars=min(char_count, MAX_EXTRACTED_CHARS),
                warnings=warnings,
                metadata={"parser": "docx"},
            )
        except Exception:
            return ExtractionResult(mime_type=DOCX_MIME, error=ExtractionError.PARSE_FAILED)
