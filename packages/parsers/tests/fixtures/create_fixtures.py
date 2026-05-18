from pathlib import Path

FIXTURE_DIR = Path(__file__).parent


def create_pdf_fixtures() -> None:
    import fitz  # type: ignore[import-untyped]

    good = fitz.open()
    for index in range(2):
        page = good.new_page()
        page.insert_text(
            (72, 72),
            f"Página {index + 1}\nLimpeza de pele custa R$120.\nAtendimento segunda 9h às 18h.\n"
            * 4,
        )
    good.save(FIXTURE_DIR / "good.pdf")

    image_only = fitz.open()
    image_only.new_page()
    image_only.save(FIXTURE_DIR / "image_only.pdf")


def create_docx_fixture() -> None:
    from docx import Document  # type: ignore[import-untyped]

    document = Document()
    document.add_heading("Serviços", level=1)
    document.add_paragraph("Limpeza de pele custa R$120 e dura cerca de 90 minutos.")
    document.add_heading("Horários", level=2)
    document.add_paragraph("Atendemos segunda das 9h às 18h.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "serviço"
    table.rows[0].cells[1].text = "preço"
    table.rows[1].cells[0].text = "peeling"
    table.rows[1].cells[1].text = "200"
    document.save(FIXTURE_DIR / "good.docx")


def create_xlsx_fixture() -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    first = workbook.active
    first.title = "servicos"
    first.append(["serviço", "preço"])
    for name, price in [("limpeza", 120), ("peeling", 200), ("hidratacao", 95)]:
        first.append([name, price])

    second = workbook.create_sheet("horarios")
    second.append(["dia", "abre", "fecha"])
    for day in ["segunda", "terça", "quarta"]:
        second.append([day, "09:00", "18:00"])
    workbook.save(FIXTURE_DIR / "good.xlsx")


def create_text_fixtures() -> None:
    (FIXTURE_DIR / "good.csv").write_text(
        "serviço,preço\nlimpeza,120\npeeling,200\nhidratacao,95\nbotox,500\nlaser,300\n",
        encoding="utf-8",
    )
    (FIXTURE_DIR / "semicolon.csv").write_text(
        "serviço;preço\nlimpeza;120\npeeling;200\n", encoding="utf-8"
    )
    (FIXTURE_DIR / "good.txt").write_text(
        "Limpeza de pele custa R$120.\n\nAtendemos segunda das 9h às 18h.\n\nAceitamos Pix.",
        encoding="utf-8",
    )
    (FIXTURE_DIR / "empty.txt").write_bytes(b"")
    (FIXTURE_DIR / "fake.pdf").write_text("isso não é um pdf", encoding="utf-8")


def main() -> None:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    create_text_fixtures()
    create_pdf_fixtures()
    create_docx_fixture()
    create_xlsx_fixture()


if __name__ == "__main__":
    main()
